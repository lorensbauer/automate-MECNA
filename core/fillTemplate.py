import pandas as pd
from docx2pdf import convert
from docxtpl import DocxTemplate
from datetime import datetime
from docx import Document
import docx2pdf


def fillTemplate(socio, plantilla_path):
    doc = Document(plantilla_path)
    replace_word = {'{{Izena}}': str(socio["Izena"]), '{{Abizenak}}': str(socio["Abizenak"]),
                    '{{NAN_zkia}}': str(socio["NAN_zkia"]), '{{Zenbatekoa}}': str(socio["Zenbatekoa"]),
                    '{{data}}': datetime.now().strftime("%Y-%m-%d")}
    # to replace words within paragraphs
    for word in replace_word:
        for p in doc.paragraphs:
            if p.text.find(word) >= 0:
                p.text = p.text.replace(word, replace_word[word])

    # to replace words within tables
    for word in replace_word:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.find(word) >= 0:
                            p.text = p.text.replace(word, replace_word[word])

    # to replace words within headers
    for word in replace_word:
        for section in doc.sections:
            header = section.header
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.find(word) >= 0:
                                p.text = p.text.replace(word, replace_word[word])

    # to replace words within footers
    for word in replace_word:
        for footer in doc.sections:
            footer = section.footer
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.find(word) >= 0:
                                p.text = p.text.replace(word, replace_word[word])
    return doc


def funcionConDocx(lista_socios, replace_word):
    doc = Document('../input/Ziurtagiria2.docx')

    lista_socios = lista_socios.to_dict(orient="records")
    for idx, socio in enumerate(lista_socios):
        # list of all words to be replaced, with its new word
        replace_word = {'{{Izena}}': str(socio["Izena"]), '{{Abizenak}}': str(socio["Abizenak"]),
                        '{{NAN_zkia}}': str(socio["NAN_zkia"]), '{{Zenbatekoa}}': str(socio["Zenbatekoa"])}

        # to replace words within paragraphs
        for word in replace_word:
            for p in doc.paragraphs:
                if p.text.find(word) >= 0:
                    p.text = p.text.replace(word, replace_word[word])

        # to replace words within tables
        for word in replace_word:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.find(word) >= 0:
                                p.text = p.text.replace(word, replace_word[word])

        # to replace words within headers
        for word in replace_word:
            for section in doc.sections:
                header = section.header
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.find(word) >= 0:
                                    p.text = p.text.replace(word, replace_word[word])

        # to replace words within footers
        for word in replace_word:
            for footer in doc.sections:
                footer = section.footer
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.find(word) >= 0:
                                    p.text = p.text.replace(word, replace_word[word])
        doc.save(f'../output/{socio["Izena"]}_MECNA.docx')


def funcionConDocxtpl(df):
    # Pasar a dict para poder reccorer
    socios = df.to_dict(orient="records")

    # Cargar plantilla
    doc = DocxTemplate("../input/Ziurtagiria1.docx")
    # doc = DocxTemplate("input/Prueba1.docx")

    # Renderizar
    doc.render({
        "socios": socios,
        "salto": "\n\n\n\n\n\n\n\n",
        "data": datetime.now().strftime("%Y-%m-%d")
    })

    # guardar resultado
    doc.save("output/Ziurtagiriak_MECNA.docx")

    return doc


if __name__ == '__main__':
    mecna_names = pd.read_excel("../input/MECNA.xlsx", usecols="A:D", header=0, index_col=None,
                                sheet_name="2025 (uztailetik)")

    mecna_names = mecna_names.rename(columns={"NAN zkia": "NAN_zkia"})
    mecna_names = mecna_names.loc[:2, :]
    doc = funcionConDocx(mecna_names)

    # Convertir a pdf
    # convert("output/Ziurtagiriak_MECNA.docx", "input/Ziurtagiria_MECNA.pdf")

    exit()
